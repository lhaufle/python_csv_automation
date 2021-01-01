import docx

class Word_Gen:

    def __init__(self, lst):
        self.user_list = lst

    #private method to be used when docs are generated
    def __create_doc(self, user_name, address, email, message):

        path = 'C:/Users/lhaufle/Desktop/MassDocs/' + user_name + '.docx'

        doc = docx.Document()
        doc.add_heading(user_name, 0)
        doc.add_paragraph(
            address, style='List Bullet'
        )
        doc.add_paragraph(
            email, style='List Bullet'
        )
        doc.add_paragraph(
            message, style='List Bullet'
        )

        doc.save(path)

    def gen_messages(self, message):
        for user in self.user_list:
            #skip the headers since it has no values
            if user['user_name'] == '':
                continue
            else:
                self.__create_doc(user['user_name'], user['address'], user['email'], message)


